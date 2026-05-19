#!/usr/bin/env python3
"""
Markdown to DOCX Converter

This script converts Markdown files to Microsoft Word DOCX format.
It supports basic Markdown syntax including headers, bold, italic, lists, and links.
"""

import argparse
import sys
import os
from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


class MarkdownToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Set up document styles for different heading levels."""
        # Configure heading styles
        styles = self.doc.styles
        
        # Heading 1
        if 'Heading 1' in styles:
            h1_style = styles['Heading 1']
            h1_style.font.size = Pt(18)
            h1_style.font.bold = True
        
        # Heading 2
        if 'Heading 2' in styles:
            h2_style = styles['Heading 2']
            h2_style.font.size = Pt(16)
            h2_style.font.bold = True
        
        # Heading 3
        if 'Heading 3' in styles:
            h3_style = styles['Heading 3']
            h3_style.font.size = Pt(14)
            h3_style.font.bold = True
    
    def convert_markdown_to_docx(self, markdown_text):
        """Convert markdown text to DOCX format."""
        lines = markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - add paragraph break
                self.doc.add_paragraph()
                i += 1
                continue
            
            # Check for headers
            if line.startswith('#'):
                self._process_header(line)
            
            # Check for unordered lists
            elif line.startswith('- ') or line.startswith('* '):
                self._process_unordered_list(lines, i)
            
            # Check for ordered lists
            elif re.match(r'^\d+\. ', line):
                self._process_ordered_list(lines, i)
            
            # Check for blockquotes
            elif line.startswith('> '):
                self._process_blockquote(lines, i)
            
            # Regular paragraph
            else:
                self._process_paragraph(line)
            
            i += 1
    
    def _process_header(self, line):
        """Process markdown header."""
        level = len(line) - len(line.lstrip('#'))
        header_text = line.lstrip('# ').strip()
        
        if level <= 3:
            heading = self.doc.add_heading(header_text, level=level)
        else:
            # For levels 4-6, use regular paragraph with bold formatting
            p = self.doc.add_paragraph()
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(12 - (level - 3))
    
    def _process_unordered_list(self, lines, start_index):
        """Process unordered list items."""
        i = start_index
        while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
            line = lines[i].strip()
            list_text = line[2:].strip()
            p = self.doc.add_paragraph(list_text, style='List Bullet')
            self._format_inline_elements(p)
            i += 1
    
    def _process_ordered_list(self, lines, start_index):
        """Process ordered list items."""
        i = start_index
        while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
            line = lines[i].strip()
            list_text = re.sub(r'^\d+\. ', '', line)
            p = self.doc.add_paragraph(list_text, style='List Number')
            self._format_inline_elements(p)
            i += 1
    
    def _process_blockquote(self, lines, start_index):
        """Process blockquote."""
        i = start_index
        quote_lines = []
        
        while i < len(lines) and lines[i].strip().startswith('> '):
            line = lines[i].strip()
            quote_text = line[2:].strip()
            quote_lines.append(quote_text)
            i += 1
        
        if quote_lines:
            quote_text = ' '.join(quote_lines)
            p = self.doc.add_paragraph()
            p.style = 'Quote'
            run = p.add_run(quote_text)
            run.italic = True
    
    def _process_paragraph(self, line):
        """Process regular paragraph."""
        if line.strip():
            p = self.doc.add_paragraph()
            self._format_inline_elements(p, line)
    
    def _format_inline_elements(self, paragraph, text=None):
        """Format inline elements like bold, italic, and links."""
        if text is None:
            # For existing paragraph, get the text from the first run
            if paragraph.runs:
                text = paragraph.runs[0].text
                paragraph.clear()
            else:
                return
        
        # Process bold text (**text** or __text__)
        text = re.sub(r'\*\*(.*?)\*\*', r'<bold>\1</bold>', text)
        text = re.sub(r'__(.*?)__', r'<bold>\1</bold>', text)
        
        # Process italic text (*text* or _text_)
        text = re.sub(r'\*(.*?)\*', r'<italic>\1</italic>', text)
        text = re.sub(r'_(.*?)_', r'<italic>\1</italic>', text)
        
        # Process inline code (`code`)
        text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)
        
        # Process links [text](url)
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<link>\1</link>', text)
        
        # Split text by tags and add runs
        parts = re.split(r'(<[^>]+>.*?</[^>]+>)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('<bold>') and part.endswith('</bold>'):
                run = paragraph.add_run(part[6:-7])
                run.bold = True
            elif part.startswith('<italic>') and part.endswith('</italic>'):
                run = paragraph.add_run(part[8:-9])
                run.italic = True
            elif part.startswith('<code>') and part.endswith('</code>'):
                run = paragraph.add_run(part[6:-7])
                run.font.name = 'Courier New'
                # Note: python-docx doesn't have built-in background color for inline code
            elif part.startswith('<link>') and part.endswith('</link>'):
                run = paragraph.add_run(part[6:-7])
                run.font.color.rgb = None  # Default link color
                run.font.underline = True
            else:
                paragraph.add_run(part)
    
    def save_docx(self, output_path):
        """Save the document to a DOCX file."""
        self.doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown files to DOCX format')
    parser.add_argument('input_file', help='Input Markdown file path')
    parser.add_argument('-o', '--output', help='Output DOCX file path (optional)')
    parser.add_argument('-v', '--verbose', action='store_true', help='Enable verbose output')
    
    args = parser.parse_args()
    
    # Check if input file exists
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Error: Input file '{args.input_file}' does not exist.")
        sys.exit(1)
    
    # Determine output file path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix('.docx')
    
    if args.verbose:
        print(f"Converting '{input_path}' to '{output_path}'...")
    
    try:
        # Read markdown file
        with open(input_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert to DOCX
        converter = MarkdownToDocxConverter()
        converter.convert_markdown_to_docx(markdown_content)
        converter.save_docx(output_path)
        
        if args.verbose:
            print(f"Successfully converted to '{output_path}'")
        else:
            print(f"Conversion complete: {output_path}")
    
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        sys.exit(1)


if __name__ == '__main__':
    main()
